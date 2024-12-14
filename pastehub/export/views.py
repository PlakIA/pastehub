from django.http import FileResponse, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404
from docx import Document

from core.storage import get_from_storage
from paste.models import Paste


def export_source(request, short_link, version):
    paste = get_object_or_404(Paste, short_link=short_link)
    paste_text = get_from_storage(f"pastes_version/{paste.id}_{version}")

    response = FileResponse(
        paste_text,
        content_type="text/plain",
        filename=f"{paste.short_link}.txt",
        as_attachment=True,
    )
    response["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response["Pragma"] = "no-cache"
    response["Expires"] = "0"

    return response


def export_json(request, short_link, version):
    paste = get_object_or_404(Paste, short_link=short_link)

    data = {
        "content": get_from_storage(f"pastes_version/{paste.id}_{version}"),
        "title": paste.title,
        "author": str(paste.author),
        "category": str(paste.category),
        "created": paste.created,
        "short_link": paste.short_link,
    }
    response = JsonResponse(
        data,
        json_dumps_params={
            "indent": 4,
            "ensure_ascii": False,
        },
    )
    response.charset = "utf-8"
    response["Content-Disposition"] = (
        f"attachment; filename={paste.short_link}.json"
    )
    response["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response["Pragma"] = "no-cache"
    response["Expires"] = "0"

    return response


def export_docx(request, short_link, version):
    paste = get_object_or_404(Paste, short_link=short_link)
    paste_text = get_from_storage(f"pastes_version/{paste.id}_{version}")

    document = Document()
    document.add_heading(paste.title, level=1)
    document.add_paragraph(f"Категория: {paste.category}")
    document.add_paragraph(f"Автор: {paste.author}")
    document.add_paragraph(f"Создана: {paste.created}")
    document.add_paragraph(paste_text)
    response = HttpResponse(
        content_type="application/vnd."
        "openxmlformats-officedocument."
        "wordprocessingml.document",
    )
    response["Content-Disposition"] = (
        f"attachment; filename={paste.short_link}.docx"
    )
    response["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response["Pragma"] = "no-cache"
    response["Expires"] = "0"
    document.save(response)

    return response


__all__ = ["export_docx", "export_json", "export_source"]
